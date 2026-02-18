"""
CalcHub Backend - FastAPI сервис для генерации норм выдачи СИЗ
"""
from fastapi import FastAPI, HTTPException, Depends, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, EmailStr
from typing import List, Optional, Dict, Any
import json
import os
import re
import io
import zipfile
import hashlib
from datetime import datetime, timedelta
import openpyxl

# Document generation
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = FastAPI(title="CalcHub API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, '..', 'src', 'data')
TEMPLATES_DIR = os.path.join(BASE_DIR, '..', 'templates')
OUTPUT_DIR = os.path.join(BASE_DIR, '..', 'outputs')
IDEN_HAZ_PATH = os.path.join(BASE_DIR, '..', 'iden_haz.xlsx')

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─── Models ───────────────────────────────────────────────────────────────────

class ProfessionEntry(BaseModel):
    profession_id: int
    profession_name: str
    hazards: List[str] = []

class GenerateRequest(BaseModel):
    professions: List[ProfessionEntry]
    doc_type: str  # "normy" or "anketa" or "all"

# ─── Data Loading ─────────────────────────────────────────────────────────────

def load_professions():
    path = os.path.join(DATA_DIR, 'professions.json')
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_hazards():
    path = os.path.join(DATA_DIR, 'hazards.json')
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_iden_haz_xlsx():
    """Load hazard SIZ data from iden_haz.xlsx"""
    wb = openpyxl.load_workbook(IDEN_HAZ_PATH)
    ws = wb.active
    
    hazards = []
    current_hazard = None
    current_event = None
    current_siz_type = None
    
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i <= 1:
            continue
        col0, col1, col2, col3, col4 = (row[j] if len(row) > j else None for j in range(5))
        
        if col0 and str(col0).strip() and not isinstance(col0, (int, float)):
            current_hazard = str(col0).strip()
            current_event = str(col1).strip() if col1 else ''
        
        if col2 and str(col2).strip() and str(col2).strip() not in ['или']:
            current_siz_type = str(col2).strip()
        
        if col3 and str(col3).strip():
            siz_name = str(col3).strip()
            norm = str(col4).strip() if col4 else ''
            
            if current_hazard:
                hazards.append({
                    'hazard': current_hazard,
                    'event': current_event or '',
                    'siz_type': current_siz_type or '',
                    'siz_name': siz_name,
                    'norm': norm
                })
    
    return hazards

# ─── SIZ Data Processing ──────────────────────────────────────────────────────

def parse_norm_string(norm_str: str):
    """Парсит строку норм выдачи"""
    if not norm_str or norm_str == 'None':
        return None, None
    
    norm_str = str(norm_str).strip()
    
    if 'износа' in norm_str.lower():
        return 'до износа', None
    
    match = re.match(r'(\d+)\s*([а-яА-Яa-zA-Z]+)', norm_str)
    if match:
        quantity = int(match.group(1))
        unit = match.group(2).strip().lower()
        
        unit_map = {
            'пара': 'пары', 'пары': 'пары',
            'шт': 'шт', 'штука': 'шт', 'штуки': 'шт',
            'комплект': 'комплекты', 'комплекты': 'комплекты',
            'мл': 'мл', 'г': 'г'
        }
        normalized_unit = unit_map.get(unit, unit)
        return quantity, normalized_unit
    
    return None, None

def extract_years_info(norm_str: str) -> str:
    """Извлекает информацию о периоде"""
    if not norm_str:
        return ''
    
    norm_lower = str(norm_str).lower()
    patterns = ['на 2 года', 'на 3 года', 'на 4 года', 'на 5 лет', 'на 1,5 года',
                'на 2 года', 'на 3 лет', 'до износа', 'определяется документами']
    
    for p in patterns:
        if p in norm_lower:
            return ''  # already has period info
    
    return ''

def format_norm_for_normy(norm_str: str) -> str:
    """Форматирует норму для документа normy-vidachi-siz"""
    if not norm_str or norm_str == 'None':
        return ''
    
    norm_str = str(norm_str).strip()
    
    # Если уже содержит период или описание - оставляем как есть
    lower = norm_str.lower()
    skip_phrases = ['на 2 год', 'на 3 год', 'на 4 год', 'на 5 лет', 'на 1,5', 
                    'износ', 'определяется', 'до']
    for phrase in skip_phrases:
        if phrase in lower:
            return norm_str
    
    # Если только цифра - добавляем "на 1 год"
    if re.match(r'^\d+$', norm_str.strip()):
        return f"{norm_str} на 1 год"
    
    # Если цифра + единица без года
    if re.match(r'^\d+\s*[а-яА-Я]+$', norm_str.strip()):
        return f"{norm_str} на 1 год"
    
    return norm_str

def get_unit_only(norm_str: str) -> str:
    """Извлекает только единицу измерения"""
    if not norm_str:
        return 'шт'
    
    lower = norm_str.lower()
    if 'пар' in lower:
        return 'пары'
    if 'комплект' in lower:
        return 'комплекты'
    if 'мл' in lower:
        return 'мл'
    return 'шт'

# ─── Get SIZ for profession ───────────────────────────────────────────────────

def get_siz_for_profession(profession_id: int) -> List[Dict]:
    """Возвращает список СИЗ для профессии из professions.json"""
    # Since we only have profession names/ids and siz_count in professions.json
    # We need the actual SIZ data - return sample data structure
    # In production this would come from the full professions.xlsx
    return []

# Load hazard SIZ data once
_hazard_siz_cache = None

def get_hazard_siz_data() -> List[Dict]:
    global _hazard_siz_cache
    if _hazard_siz_cache is None:
        _hazard_siz_cache = load_iden_haz_xlsx()
    return _hazard_siz_cache

# ─── Document Generation ──────────────────────────────────────────────────────

def add_cell_borders(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '4')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), '000000')
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def set_cell_text(cell, text, font_size=10, bold=False):
    """Sets cell text with formatting"""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.clear()
    run = para.add_run(str(text) if text else '')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    add_cell_borders(cell)

def generate_normy_document(professions_list: List[ProfessionEntry]) -> bytes:
    """Генерирует документ норм выдачи СИЗ"""
    template_path = os.path.join(TEMPLATES_DIR, 'normy-vidachi-siz.docx')
    
    try:
        doc = Document(template_path)
    except Exception:
        doc = Document()
        doc.add_heading('Нормы выдачи СИЗ', 0)
    
    hazard_siz = get_hazard_siz_data()
    
    # Build rows data
    rows_data = []
    row_num = 1
    
    for entry in professions_list:
        prof_rows_count = 0
        
        # Base SIZ from profession (placeholder - in production from xlsx)
        base_siz = [
            {'type': 'Спецодежда', 'name': f'СИЗ для {entry.profession_name}', 
             'norm': '1 шт', 'source': 'professions'}
        ]
        
        for siz in base_siz:
            norm_formatted = format_norm_for_normy(siz['norm'])
            rows_data.append({
                'num': row_num if prof_rows_count == 0 else '',
                'profession': entry.profession_name if prof_rows_count == 0 else '',
                'type': siz['type'],
                'name': siz['name'],
                'norm': norm_formatted,
                'basis': f"{row_num}, Приложение № 1 к приказу Минтруда России от 29.10.2021 №767н"
            })
            prof_rows_count += 1
        
        # SIZ from hazards
        for hazard_name in entry.hazards:
            matching = [h for h in hazard_siz if h['hazard'] == hazard_name]
            for h in matching[:3]:  # limit to first 3 matches per hazard
                norm_formatted = format_norm_for_normy(h['norm'])
                rows_data.append({
                    'num': row_num if prof_rows_count == 0 else '',
                    'profession': entry.profession_name if prof_rows_count == 0 else '',
                    'type': h['siz_type'],
                    'name': h['siz_name'],
                    'norm': norm_formatted,
                    'basis': f"{row_num}, Приложение № 2 к приказу Минтруда России от 29.10.2021 №767н"
                })
                prof_rows_count += 1
        
        row_num += 1
    
    # Find or create table
    target_table = None
    for table in doc.tables:
        if len(table.columns) >= 6:
            target_table = table
            break
    
    if not target_table:
        # Create new table
        doc.add_paragraph()
        target_table = doc.add_table(rows=1, cols=6)
        target_table.style = 'Table Grid'
        header_row = target_table.rows[0]
        headers = ['№ п.п', 'Наименование профессии (должности)', 'Тип СИЗ', 
                   'Наименование СИЗ', 'Нормы выдачи', 'Основание выдачи СИЗ']
        for i, h in enumerate(headers):
            set_cell_text(header_row.cells[i], h, bold=True)
    
    # Add data rows
    for row_data in rows_data:
        new_row = target_table.add_row()
        cells = new_row.cells
        set_cell_text(cells[0], row_data['num'])
        set_cell_text(cells[1], row_data['profession'])
        set_cell_text(cells[2], row_data['type'])
        set_cell_text(cells[3], row_data['name'])
        set_cell_text(cells[4], row_data['norm'])
        set_cell_text(cells[5], row_data['basis'])
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def generate_anketa_document(entry: ProfessionEntry) -> bytes:
    """Генерирует личную карточку СИЗ для профессии"""
    template_path = os.path.join(TEMPLATES_DIR, 'personal_anketa.docx')
    
    try:
        doc = Document(template_path)
    except Exception:
        doc = Document()
        doc.add_heading(f'Личная карточка учёта выдачи СИЗ', 0)
        doc.add_paragraph(f'Профессия: {entry.profession_name}')
    
    hazard_siz = get_hazard_siz_data()
    
    # Build SIZ list
    siz_items = []
    
    # Base items
    siz_items.append({
        'name': f'СИЗ для {entry.profession_name}',
        'punkt': '1',
        'unit': 'шт',
        'qty': '1 на 1 год'
    })
    
    # Hazard items
    for hazard_name in entry.hazards:
        matching = [h for h in hazard_siz if h['hazard'] == hazard_name]
        for h in matching[:3]:
            unit = get_unit_only(h['norm'])
            qty = format_norm_for_normy(h['norm'])
            siz_items.append({
                'name': h['siz_name'],
                'punkt': '№' + h.get('pp', '1'),
                'unit': unit,
                'qty': qty
            })
    
    # Find SIZ table
    target_table = None
    for table in doc.tables:
        first_cell_text = table.rows[0].cells[0].text.strip() if table.rows else ''
        if 'Наименование СИЗ' in first_cell_text or len(table.columns) >= 4:
            target_table = table
            break
    
    if not target_table:
        doc.add_paragraph()
        target_table = doc.add_table(rows=1, cols=4)
        target_table.style = 'Table Grid'
        header_row = target_table.rows[0]
        headers = ['Наименование СИЗ', 'Пункт норм', 'Единица измерения', 'Количество на период']
        for i, h in enumerate(headers):
            set_cell_text(header_row.cells[i], h, bold=True)
    
    for item in siz_items:
        new_row = target_table.add_row()
        cells = new_row.cells
        set_cell_text(cells[0], item['name'])
        set_cell_text(cells[1], item['punkt'])
        set_cell_text(cells[2], item['unit'])
        set_cell_text(cells[3], item['qty'])
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

# ─── Routes ───────────────────────────────────────────────────────────────────

@app.get("/api/professions")
def get_professions(search: str = '', limit: int = 50):
    professions = load_professions()
    if search:
        search_lower = search.lower()
        professions = [p for p in professions if search_lower in p['name'].lower()]
    return professions[:limit]

@app.get("/api/hazards")
def get_hazards():
    hazards = load_hazards()
    # Return unique hazard names
    seen = set()
    unique = []
    for h in hazards:
        if h['hazard'] not in seen:
            seen.add(h['hazard'])
            unique.append({'name': h['hazard'], 'event': h.get('event', '')})
    return unique

@app.post("/api/documents/generate-normy")
def generate_normy(request: GenerateRequest):
    if not request.professions:
        raise HTTPException(400, "Список профессий пуст")
    
    doc_bytes = generate_normy_document(request.professions)
    
    filename = f"normy-vydachi-siz-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
    
    return Response(
        content=doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )

@app.post("/api/documents/generate-anketa")
def generate_anketa(request: GenerateRequest):
    if not request.professions:
        raise HTTPException(400, "Список профессий пуст")
    
    if len(request.professions) == 1:
        doc_bytes = generate_anketa_document(request.professions[0])
        prof_name = re.sub(r'[<>:"/\\|?*]', '', request.professions[0].profession_name)
        filename = f"anketa-{prof_name[:30]}.docx"
        
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
        )
    
    # Multiple professions - return zip
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for entry in request.professions:
            doc_bytes = generate_anketa_document(entry)
            prof_name = re.sub(r'[<>:"/\\|?*]', '', entry.profession_name)
            filename = f"anketa-{prof_name[:30]}.docx"
            zf.writestr(filename, doc_bytes)
    
    zip_buffer.seek(0)
    return Response(
        content=zip_buffer.read(),
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=ankety-siz.zip"}
    )

@app.get("/api/health")
def health():
    return {"status": "ok", "version": "1.0.0"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
